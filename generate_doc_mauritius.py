"""
Generate AstraFlow System Design & Architecture document
with Mauritius fuel & power sector focus.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, cells_text, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        if header:
            run.bold = True
            set_cell_shading(cell, "003087")
            run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            run.font.color.rgb = RGBColor(26, 28, 30)
    return row


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    for level, size, color in [
        ("Heading 1", 20, "003087"),
        ("Heading 2", 15, "003087"),
        ("Heading 3", 12, "1c4197"),
    ]:
        s = doc.styles[level]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(16 if level == "Heading 1" else 10)
        s.paragraph_format.space_after = Pt(6)

    # ═══════════════════ TITLE PAGE ════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AstraFlow")
    run.font.size = Pt(44)
    run.font.color.rgb = RGBColor.from_string("003087")
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("System Design & Algorithmic Architecture")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("1c4197")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mauritius Fuel & Power Sector Edition")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("d32f2f")
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Intelligent Fuel Insights for the Mauritian Economy")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("747683")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Document Version 2.0  |  {datetime.date.today().strftime('%B %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("444652")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Comprehensive System Review — Focused on Mauritius Fuel Imports, Power Generation & Energy Transition")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("747683")

    doc.add_page_break()

    # ═══════════════════ TABLE OF CONTENTS ═════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc = [
        ("1.", "Executive Summary", 3),
        ("2.", "System Architecture Overview", 4),
        ("3.", "Frontend Architecture", 5),
        ("4.", "Backend Architecture", 7),
        ("5.", "Algorithms & Processing Logic", 9),
        ("5.1", "JWT Authentication Flow", 9),
        ("5.2", "Dashboard Personalisation — Mauritius Sector Mapping", 10),
        ("5.3", "Consumption Prediction Engine", 11),
        ("5.4", "Price Alert Threshold", 11),
        ("5.5", "Fuel Price History Visualisation", 12),
        ("5.6", "Survey Multi-Step Form", 12),
        ("5.7", "Protected Route Guard", 13),
        ("5.8", "Graceful Degradation & Mock Fallback", 13),
        ("5.9", "Survey Insights Aggregation", 14),
        ("6.", "Data Flow Analysis", 15),
        ("7.", "Security Analysis", 17),
        ("8.", "Performance Analysis", 19),
        ("9.", "System Integration — Collective Behaviour", 21),
        ("10.", "Deployment Architecture", 22),
        ("11.", "Mauritius Energy Policy Context", 23),
        ("12.", "Future Roadmap", 24),
    ]
    for num, title, _ in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        indent = "    " if "." in num and not num.endswith(".") else ""
        run = p.add_run(f"{indent}{num}  {title}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("1a1c1e")

    doc.add_page_break()

    # ═══════════════════ 1. EXECUTIVE SUMMARY ═════════════════════════════
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "AstraFlow is an AI-Based Fuel Price Forecasting and Evaluation System purpose-built for "
        "the Republic of Mauritius — a Small Island Developing State where 90.9% of primary energy "
        "requirements are sourced from imports (61.1% petroleum products, 29.8% coal as of 2024). "
        "Fuel imports account for 21.2% of Mauritius' total import bill, making energy security a "
        "critical macroeconomic concern. The system is designed as a full-stack mobile-first application "
        "using Expo (React Native) for the frontend and FastAPI (Python) for the backend, backed by "
        "a PostgreSQL database."
    )
    doc.add_paragraph(
        "AstraFlow serves the full spectrum of Mauritius' energy economy — from the Central Electricity "
        "Board (CEB) and Independent Power Producers managing heavy fuel oil and coal procurement, "
        "to the sugar industry leveraging bagasse cogeneration, to transport and logistics fleet operators, "
        "manufacturing enterprises in the Export Processing Zone, and the tourism and hospitality sector. "
        "The application provides real-time fuel price tracking aligned with the State Trading Corporation's "
        "Petroleum Pricing Mechanism (PPM), consumption cost prediction with carbon footprint estimation "
        "benchmarked against Mauritius' Nationally Determined Contributions, sector-specific impact assessment, "
        "personalised recommendations, and fuel impact surveys."
    )
    doc.add_paragraph(
        "The system uses JWT-based authentication, a token-driven Azure Clarity design system, graceful "
        "degradation with mock data fallback, and a clearly separated service-route-model architecture on "
        "the backend. This document provides a comprehensive review of the system design, every algorithm "
        "driving the application's features, and an analysis of how all components integrate as a collective "
        "system — with the entire context grounded in Mauritius' fuel and power sector realities."
    )

    # ═══════════════════ 2. SYSTEM ARCHITECTURE ══════════════════════════
    doc.add_heading("2. System Architecture Overview", level=1)

    doc.add_heading("2.1 High-Level Architecture", level=2)
    doc.add_paragraph(
        "AstraFlow follows a modern two-tier monorepo architecture. The frontend and backend coexist in a "
        "single repository, sharing database schemas and design references. Communication occurs exclusively "
        "over HTTP REST APIs using JSON as the data interchange format."
    )

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(["Layer", "Technology", "Version"]):
        hdr.cells[i].text = text
        set_cell_shading(hdr.cells[i], "003087")
        run = hdr.cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    layers = [
        ("Mobile / Web Framework", "Expo (React Native)", "~56.0.9"),
        ("Routing", "Expo Router (file-based)", "~56.2.9"),
        ("UI Component Library", "Gluestack UI", "~1.1.73"),
        ("Styling Engine", "NativeWind (Tailwind for RN)", "~4.2.5"),
        ("Animations", "React Native Reanimated", "~4.3.1"),
        ("Backend Framework", "FastAPI (Python)", "0.115.0"),
        ("ASGI Server", "Uvicorn", "0.30.0"),
        ("ORM", "SQLAlchemy", "2.0.50"),
        ("Database", "PostgreSQL (Supabase or on-premise)", "—"),
        ("Migrations", "Alembic", "1.18.4"),
        ("JWT / Auth", "python-jose + passlib (bcrypt)", "3.5.0 / 1.7.4"),
        ("ML / Analysis", "scikit-learn, pandas, numpy", "1.5.0 / 2.2.0 / 1.26.0"),
        ("Deployment", "Render / Vercel or on-premise (CEB/Ministry of Energy)", "—"),
    ]
    for layer in layers:
        add_table_row(table, layer)

    doc.add_paragraph()

    doc.add_heading("2.2 Monorepo Structure", level=2)
    doc.add_paragraph(
        "The repository is organised into the following top-level directories:"
    )
    for item in [
        "src/app/ — Expo Router file-based routes (login, register, tabs dashboard, predict, prices, profile, survey).",
        "src/services/ — API client and auth token management.",
        "src/theme/ — Azure Clarity Gluestack UI configuration.",
        "backend/ — FastAPI application (routes/, models/, services/, db/).",
        "backend/routes/ — HTTP endpoints (auth, dashboard, predict, surveys).",
        "backend/models/ — SQLAlchemy ORM models (User, FuelPrice, Recommendation, Survey).",
        "backend/services/ — Business logic (JWT auth with bcrypt).",
        "backend/ml/ — Placeholder for future ML models (scikit-learn, pandas, numpy ready).",
        "database/ — SQL schema with Mauritius-specific seed data.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(1)

    doc.add_page_break()

    # ═══════════════════ 3. FRONTEND ARCHITECTURE ═════════════════════════
    doc.add_heading("3. Frontend Architecture", level=1)

    doc.add_heading("3.1 File-Based Routing (Expo Router)", level=2)
    doc.add_paragraph(
        "AstraFlow uses Expo Router's file-based routing. Every file under src/app/ "
        "automatically becomes a route. The (tabs) group creates a bottom tab navigator."
    )
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(["Route", "File", "Purpose"]):
        hdr.cells[i].text = text
        set_cell_shading(hdr.cells[i], "003087")
        run = hdr.cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    for row_data in [
        ("/", "_layout.tsx", "Root layout, splash screen, auth guard, GluestackUI provider"),
        ("/login", "login.tsx", "Email/password login with validation"),
        ("/register", "register.tsx", "Registration with Mauritius sector chip selector"),
        ("/ (tabs index)", "(tabs)/index.tsx", "Dashboard — Mogas/Gas Oil prices, STC trends, sector risk"),
        ("/prices", "(tabs)/prices.tsx", "Price history referencing PPM adjustment dates"),
        ("/predict", "(tabs)/predict.tsx", "Consumption cost (MUR) and carbon footprint calculator"),
        ("/profile", "(tabs)/profile.tsx", "User account info and logout"),
        ("/survey", "survey.tsx", "4-step multi-part fuel impact survey form"),
    ]:
        add_table_row(table, row_data)
    doc.add_paragraph()

    doc.add_heading("3.2 Root Layout & Protected Route Guard", level=2)
    doc.add_paragraph(
        "The root _layout.tsx manages a branded splash screen (AstraFlow logo with 'Intelligent Fuel Insights' "
        "tagline) lasting 1500ms, and a useProtectedRoute() hook that reads the current route segment via "
        "useSegments() and checks isAuthenticated() from localStorage. Unauthenticated users are redirected "
        "to /login; authenticated users on auth screens are redirected to the tabs dashboard."
    )

    doc.add_heading("3.3 API Client Layer", level=2)
    doc.add_paragraph(
        "The API client (src/services/api.ts) wraps fetch into a typed abstraction with: auto-injection of "
        "Bearer JWT tokens, automatic 401 handling (clear token + redirect to login), a custom ApiError class "
        "with status and detail fields, and typed namespaces (api.auth, api.dashboard, api.predict, api.surveys) "
        "providing compile-time type safety."
    )

    doc.add_heading("3.4 Auth Service (Frontend)", level=2)
    doc.add_paragraph(
        "Token and user data persist in localStorage under astraflow_token and astraflow_user. "
        "The service provides getToken(), setToken(), clearToken(), getUser(), setUser(), and isAuthenticated(). "
        "Storage access is platform-aware — returns null on native mobile where secure storage would be used."
    )

    doc.add_heading("3.5 Azure Clarity Design System", level=2)
    doc.add_paragraph(
        "The Azure Clarity design system defines the application's visual language with principles of "
        "'high-clarity minimalism' and 'trustworthy professionalism.' Key tokens include:"
    )
    for item in [
        "Primary colour: Deep Blue (#003087), referencing Mauritius' Indian Ocean identity.",
        "Fuel-specific colours: price-increase (#d32f2f), price-decrease (#2e7d32).",
        "Typography: Work Sans (headings), Inter (body), monospace (code).",
        "Spacing: 4px baseline grid; radii: sm=4px, md=8px, lg=12px, xl=16px.",
        "Tonal layering: surface colour variants instead of heavy shadows.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ═══════════════════ 4. BACKEND ARCHITECTURE ══════════════════════════
    doc.add_heading("4. Backend Architecture", level=1)

    doc.add_heading("4.1 Application Entry & Lifespan", level=2)
    doc.add_paragraph(
        "FastAPI initialises with an async lifespan that calls init_db() on startup, auto-creating "
        "database tables via SQLAlchemy's Base.metadata.create_all(). CORS allows all origins during "
        "development. Four routers are registered: auth, dashboard, predict, surveys, plus /api/health."
    )

    doc.add_heading("4.2 Database Layer", level=2)
    doc.add_paragraph(
        "SQLAlchemy engine connects to PostgreSQL (Supabase or on-premise within CEB/Ministry of Energy "
        "infrastructure). Connection pool: pool_size=5, max_overflow=10, pool_pre_ping=True for connection "
        "health checks. The get_db() generator is used as a FastAPI dependency."
    )

    doc.add_heading("4.3 ORM Models — Mauritius Sector Edition", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(["Model", "Table", "Key Fields"]):
        hdr.cells[i].text = text
        set_cell_shading(hdr.cells[i], "003087")
        run = hdr.cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    for row_data in [
        ("User", "users", "id (UUID), email, password_hash, full_name, sector (Enum: power_gen, sugar, transport, manufacturing, tourism)"),
        ("FuelPrice", "fuel_prices", "id, date, fuel_type (mogas_95, gas_oil, hfo, kerosene, coal, bagasse), price (MUR)"),
        ("Recommendation", "recommendations", "id, sector, content, risk_level, valid_from, valid_to"),
        ("Survey", "surveys", "id, user_id (FK), sector, monthly_fuel_spend (MUR), impact_level, concern_areas (ARRAY)"),
    ]:
        add_table_row(table, row_data)
    doc.add_paragraph()

    doc.add_heading("4.4 Route Modules — Mauritius Context", level=2)

    doc.add_heading("Auth Routes", level=3)
    doc.add_paragraph(
        "POST /api/auth/register validates the sector against the Mauritius Sector enum, checks for "
        "duplicate email, creates a User with bcrypt-hashed password, generates a JWT (30-day expiry, "
        "HS256), and returns the token with user metadata."
    )
    doc.add_paragraph(
        "POST /api/auth/login verifies email/password against the bcrypt hash and returns a JWT."
    )

    doc.add_heading("Dashboard Routes — Mauritius Power Sector", level=3)
    doc.add_paragraph(
        "GET /api/dashboard is protected by get_current_user(). It returns current Mogas and Gas Oil "
        "prices as set by the STC's Petroleum Pricing Committee (PPC), sector-specific recommendations "
        "mapped to Mauritius business sectors, risk levels, and a market update referencing the Price "
        "Stabilisation Account and global Brent crude trends."
    )

    doc.add_heading("Predict Routes — MUR-Based Calculations", level=3)
    doc.add_paragraph(
        "GET /api/predict?liters=X computes total cost in Mauritian Rupees, carbon footprint in kg CO₂ "
        "benchmarked against Mauritius' 4.3M ton reduction target, and future projected loss using "
        "sector-specific price trends."
    )

    doc.add_heading("Survey Routes", level=3)
    doc.add_paragraph(
        "POST /api/surveys submits fuel impact surveys. GET /api/surveys lists user responses. "
        "GET /api/surveys/insights returns aggregate statistics including impact distribution across "
        "Mauritius sectors."
    )

    doc.add_heading("4.5 Auth Service (Backend)", level=2)
    doc.add_paragraph(
        "bcrypt password hashing via passlib CryptContext. JWT creation with HS256, 30-day expiry, "
        "user UUID as sub claim. Token decoding verifies signature and expiration, returning the user "
        "ID on success or None on any JWTError."
    )

    doc.add_heading("4.6 Seed Data — Mauritius Market", level=2)
    doc.add_paragraph(
        "The seed script populates fuel prices using actual STC retail data (Mogas Rs 64.25/L, "
        "Gas Oil Rs 71.25/L as of April 2026) plus Heavy Fuel Oil prices for power generation. "
        "Five recommendations are seeded — one per Mauritius sector — reflecting current PPC "
        "adjustment thresholds and the Price Stabilisation Account balance."
    )

    doc.add_page_break()

    # ═══════════════════ 5. ALGORITHMS ════════════════════════════════════
    doc.add_heading("5. Algorithms & Processing Logic", level=1)
    doc.add_paragraph(
        "This section describes every distinct algorithm implemented in AstraFlow, with all context "
        "grounded in Mauritius' fuel and power sector realities."
    )

    doc.add_heading("5.1 JWT Authentication Flow", level=2)
    doc.add_paragraph(
        "The authentication pipeline follows five phases: (1) Registration with bcrypt hashing, sector "
        "validation against the Mauritius Sector enum, and User record creation. (2) Login with password "
        "verification and JWT issuance. (3) Token storage in localStorage. (4) Automatic Bearer token "
        "header injection on every API call. (5) Token verification via get_current_user() which extracts, "
        "decodes, and validates the token against the database on every protected route."
    )

    doc.add_heading("5.2 Dashboard Personalisation — Mauritius Sector Mapping", level=2)
    doc.add_paragraph(
        "The dashboard uses a RECOMMENDATIONS dictionary mapping five Mauritius-focused sectors to "
        "personalised advice, risk levels, and impact assessments. Each sector receives contextually "
        "relevant guidance aligned with Mauritius' energy landscape."
    )
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(["Mauritius Sector", "Risk", "Impact", "Recommendation Theme"]):
        hdr.cells[i].text = text
        set_cell_shading(hdr.cells[i], "003087")
        run = hdr.cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    for row_data in [
        ("Power Generation (CEB, IPPs, CT Power)", "High", "High",
         "HFO/coal procurement hedging, IPP tariff negotiation, fuel switching to LNG/gas"),
        ("Sugar & Bagasse (Sugar Mills)", "Moderate", "High",
         "Bagasse cogeneration optimisation, off-season fuel efficiency, MARENA biomass incentives"),
        ("Transport & Logistics (Fleet Operators)", "High", "High",
         "Fleet electrification, route optimisation via CEB EV tariff, STC bulk fuel procurement"),
        ("Manufacturing (EPZ, Textiles, Seafood)", "Low", "Low",
         "Energy audit subsidies, solar PV adoption, CEB industrial tariff analysis"),
        ("Tourism & Hospitality (Hotels, Resorts)", "Moderate", "Medium",
         "Energy efficiency certifications, solar water heating, backup generation fuel strategy"),
    ]:
        add_table_row(table, row_data)
    doc.add_paragraph()

    doc.add_heading("5.3 Consumption Prediction Engine — MUR Basis", level=2)
    doc.add_paragraph(
        "The prediction algorithm computes financial and environmental impact in Mauritian Rupees:"
    )
    doc.add_paragraph(
        "Total Cost (MUR): total_cost = litres x price_per_litre. The price per litre reflects the "
        "current STC retail price for the selected fuel type (Mogas at Rs 64.25/L or Gas Oil at Rs 71.25/L "
        "as of April 2026, adjusted at each PPC meeting).",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Carbon Footprint: carbon_footprint_kg = litres x 2.3 kg CO2. This aligns with the national "
        "target of reducing CO2 output by 4.3 million tons under Mauritius' Nationally Determined "
        "Contributions to the Paris Agreement.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Future Projected Loss: future_loss = total_cost x (future_increase_pct / 100). The "
        "future_increase_pct reflects the PPM's 4% automatic adjustment threshold — the point at which "
        "the PPC typically revises retail prices based on the 6-month Platts average.",
        style="List Bullet"
    )

    doc.add_heading("5.4 Price Alert Threshold — PPM-Based", level=2)
    doc.add_paragraph(
        "The price alert algorithm triggers when consumption exceeds 40 litres (approximately one "
        "month of average fuel use for a small commercial vehicle in Mauritius). The alert message "
        "includes the projected cost increase based on the PPM's 4% adjustment threshold and references "
        "the current Price Stabilisation Account balance."
    )

    doc.add_heading("5.5 Fuel Price History Visualisation — STC Data", level=2)
    doc.add_paragraph(
        "The price history screen renders a 7-day horizontal bar chart using native React Native views. "
        "Mogas bars use Deep Blue (#003087) and Gas Oil bars use red (#d32f2f). Bar width is proportional "
        "to (day.price / maxPrice) x 100%, reflecting the actual STC retail price trajectory. The data "
        "aligns with the historical price registry maintained by the State Trading Corporation."
    )

    doc.add_heading("5.6 Survey Multi-Step Form", level=2)
    doc.add_paragraph(
        "A 4-step form collects: (1) Monthly fuel spend in MUR, (2) Impact level (Low, Medium, High, "
        "Severe), (3) Concern areas (Fuel import costs, CEB tariff fluctuations, PPM adjustment frequency, "
        "Renewable transition costs, Supply chain logistics), (4) Comments. On submission, data is posted "
        "to POST /api/surveys and stored against the user's sector."
    )

    doc.add_heading("5.7 Protected Route Guard", level=2)
    doc.add_paragraph(
        "The useProtectedRoute() hook reads the route segment, checks isAuthenticated(), and redirects: "
        "unauthenticated users to /login, authenticated users on auth screens to / (tabs). Runs once "
        "via a checked flag to prevent infinite loops."
    )

    doc.add_heading("5.8 Graceful Degradation & Mock Fallback", level=2)
    doc.add_paragraph(
        "When the backend is unreachable, the dashboard falls back to MOCK_DATA (STC prices and "
        "sector recommendations computed client-side). The predictor falls back to MOCK_RESULT which "
        "computes the same MUR-based formulas using identical constants. Skeleton loading provides "
        "placeholder UI during the loading phase."
    )

    doc.add_heading("5.9 Survey Insights Aggregation", level=2)
    doc.add_paragraph(
        "The GET /api/surveys/insights endpoint aggregates survey data across Mauritius sectors: "
        "impact distribution (count by sector and severity), total survey count, and the authenticated "
        "user's 5 most recent submissions. This provides CEB and MARENA policymakers with visibility "
        "into sector-level fuel cost impacts."
    )

    doc.add_page_break()

    # ═══════════════════ 6. DATA FLOW ════════════════════════════════════
    doc.add_heading("6. Data Flow Analysis — Mauritius Use Cases", level=1)
    doc.add_paragraph(
        "This section traces end-to-end data movement for each major user journey within the "
        "Mauritius power sector context."
    )
    journeys = [
        ("Power Plant Fuel Procurement Journey",
         "CEB/IPP procurement officer logs in -> Dashboard shows current HFO price (Rs 42.00/L) and "
         "Mogas/Gas Oil retail prices -> Predict tool calculates cost of 50,000L HFO delivery = "
         "Rs 2,100,000 -> Price alert triggers if above threshold -> Recommendation advises "
         "fuel hedging based on 6-month Platts average and PSA balance."),
        ("Sugar Mill Cogeneration Optimisation Journey",
         "Sugar mill operations manager registers under 'Sugar & Bagasse' sector -> Dashboard shows "
         "bagasse cogeneration efficiency metrics and HFO backup costs -> Predict tool models "
         "fuel switching scenarios (bagasse vs HFO) -> Survey captures seasonal fuel spend patterns -> "
         "Insights feed into MARENA biomass incentive programme analysis."),
        ("Fleet Operator Fuel Management Journey",
         "Transport company manager logs in -> Dashboard shows Mogas/Gas Oil STC retail prices and "
         "trends -> Predict tool estimates monthly fleet fuel cost for 40 vehicles (approx 4,000L/month) "
         "= Rs 285,000 at current Gas Oil price -> Price alert warns if PPM triggers adjustment -> "
         "Recommendation suggests bulk procurement through STC and EV fleet transition."),
        ("Manufacturing Energy Audit Journey",
         "Textile factory energy manager registers -> Dashboard shows industrial HFO pricing and "
         "CEB electricity tariff trends -> Survey captures monthly fuel spend and impact areas -> "
         "Insights correlate fuel costs with manufacturing output -> Recommendation promotes "
         "solar PV adoption and energy audit subsidies."),
        ("Tourism Energy Efficiency Journey",
         "Hotel facilities manager logs in -> Dashboard shows Mogas for backup generators and "
         "Gas Oil for heating -> Predict tool calculates carbon footprint against 4.3M ton national "
         "target -> Recommendation advises solar water heating and energy efficiency certifications."),
    ]
    for title, flow_text in journeys:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph(flow_text)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ═══════════════════ 7. SECURITY ═════════════════════════════════════
    doc.add_heading("7. Security Analysis", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(["Aspect", "Implementation", "Assessment"]):
        hdr.cells[i].text = text
        set_cell_shading(hdr.cells[i], "003087")
        run = hdr.cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    for row_data in [
        ("Password hashing", "bcrypt via passlib CryptContext", "Strong — recommended for government-grade applications"),
        ("Token signing", "HS256 (HMAC-SHA256)", "Adequate for single-service; RS256 recommended for distributed CEB/Ministry deployments"),
        ("Token expiry", "30 days", "Long-lived; consider refresh tokens for production government use"),
        ("Secret key", "Env var SECRET_KEY", "Dev fallback must be changed for production"),
        ("Token storage", "localStorage (web)", "Acceptable for web; native should use Secure Enclave/Keychain"),
    ]:
        add_table_row(table, row_data)
    doc.add_paragraph()

    doc.add_heading("7.1 Mauritius-Specific Considerations", level=2)
    doc.add_paragraph(
        "Data residency: For deployment within CEB or Ministry of Energy infrastructure, all user and "
        "fuel price data would remain within Mauritius' jurisdiction, compliant with the Data Protection "
        "Act 2017 administered by the Data Protection Office.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Procurement security: Integration with STC's Petroleum Pricing Mechanism data requires "
        "secure API authentication and audit logging per public procurement standards.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Critical infrastructure: If deployed for CEB operational use, the system should undergo "
        "National Computer Board security assessment and comply with the National Cybersecurity Strategy.",
        style="List Bullet"
    )

    doc.add_heading("7.2 Identified Security Gaps", level=2)
    for item in [
        "No rate limiting on auth endpoints — susceptible to brute-force attacks.",
        "No server-side password strength enforcement beyond frontend validation.",
        "No token revocation or blacklist mechanism — leaked JWT valid for 30 days.",
        "CORS currently open to all origins (acceptable for development only).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ═══════════════════ 8. PERFORMANCE ══════════════════════════════════
    doc.add_heading("8. Performance Analysis", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(["Area", "Current Behaviour", "Assessment"]):
        hdr.cells[i].text = text
        set_cell_shading(hdr.cells[i], "003087")
        run = hdr.cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    for row_data in [
        ("Initial load", "Splash (1500ms) + auth + API fetch", "Sequential; could parallelise"),
        ("Skeleton loading", "Placeholder blocks while loading", "Good UX — reduces perceived wait"),
        ("Connection pooling", "pool_size=5, max_overflow=10", "Adequate for CEB/Ministry departmental use"),
        ("Survey insights", "Full table scan for impact count", "Degrades with scale; use GROUP BY"),
        ("Predict endpoint", "O(1) — pure arithmetic", "Fastest endpoint, no DB reads needed"),
    ]:
        add_table_row(table, row_data)
    doc.add_paragraph()

    doc.add_heading("8.1 Optimisation Opportunities", level=2)
    for item in [
        "Replace survey impact distribution with GROUP BY aggregation query.",
        "Implement lightweight client-side caching (React Query / SWR) for dashboard data.",
        "Lazy-load screens outside the initial route to reduce bundle size.",
        "Combine splash display with parallel API fetch to reduce startup latency.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ═══════════════════ 9. SYSTEM INTEGRATION ════════════════════════════
    doc.add_heading("9. System Integration — Collective Behaviour", level=1)

    doc.add_heading("9.1 The Request-Response Cycle", level=2)
    doc.add_paragraph(
        "Every authenticated interaction follows: User action -> Screen state update -> API dispatch "
        "with Bearer token -> FastAPI routing -> get_current_user() dependency -> Route logic -> "
        "SQLAlchemy query -> PostgreSQL response -> JSON return -> setState() -> Re-render. "
        "Failure at any step triggers the error handling path (error banner, mock data fallback, "
        "and retry logic)."
    )

    doc.add_heading("9.2 Component Dependency Graph", level=2)
    for item in [
        "_layout.tsx depends on: theme/azure-clarity.ts, services/auth.ts, expo-router.",
        "(tabs)/index.tsx depends on: services/api.ts, services/auth.ts, skeleton components.",
        "(tabs)/predict.tsx depends on: services/api.ts (api.predict.get).",
        "survey.tsx depends on: services/api.ts (api.surveys.submit).",
        "services/api.ts depends on: services/auth.ts (getToken, clearToken).",
        "backend/routes/dashboard.py depends on: auth service, User model, database module.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9.3 Failure Mode Analysis — Mauritius Context", level=2)
    failures = [
        ("Backend unreachable", "Frontend falls back to mock STC prices. Error banner displayed. All UI remains functional. Critical for field use where CEB connectivity may be intermittent."),
        ("JWT expired", "API client clears token. Route guard redirects to /login. User re-authenticates. Acceptable for non-critical advisory use."),
        ("Database connection failure", "Startup init_db() throws. Requires restart or health check. For production, implement DB connection retry and failover to secondary replica within Mauritius."),
        ("Invalid sector registration", "Server returns 400. User corrects selection. Prevents data contamination across Mauritius sectors."),
    ]
    for title, desc in failures:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(desc).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ═══════════════════ 10. DEPLOYMENT ══════════════════════════════════
    doc.add_heading("10. Deployment Architecture", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(["Component", "Platform", "Configuration"]):
        hdr.cells[i].text = text
        set_cell_shading(hdr.cells[i], "003087")
        run = hdr.cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    for row_data in [
        ("Backend API", "Render / Ministry of Energy on-premise", "Python 3.12 Docker container. On-premise: CentOS/Ubuntu server within CEB datacentre."),
        ("Frontend Web", "Vercel / CEB internal network", "Static export via npx expo export --platform web. Internal deployment via government intranet."),
        ("Database", "Supabase / PostgreSQL on-premise", "Managed Supabase for SaaS. Local PostgreSQL for government deployment with daily backups."),
        ("Containerisation", "Docker", "Python 3.12-slim base. Reproducible builds for audit compliance."),
    ]:
        add_table_row(table, row_data)
    doc.add_paragraph()

    doc.add_heading("10.1 Environment Configuration", level=2)
    doc.add_paragraph(
        "The backend reads DATABASE_URL and SECRET_KEY from environment variables. For Mauritius "
        "government deployment, these would be managed through the Ministry of Energy's secure "
        "configuration management system, with secrets rotated per National Cybersecurity Strategy guidelines."
    )

    doc.add_page_break()

    # ═══════════════════ 11. MAURITIUS ENERGY POLICY CONTEXT ═════════════
    doc.add_heading("11. Mauritius Energy Policy Context", level=1)

    doc.add_paragraph(
        "AstraFlow operates within the framework of Mauritius' national energy policies and targets. "
        "This section contextualises the system's algorithms and data against the country's energy "
        "landscape."
    )

    doc.add_heading("11.1 Energy Mix & Import Dependence", level=2)
    doc.add_paragraph(
        "Mauritius relies on imports for 90.9% of its primary energy requirements. In 2024, petroleum "
        "products accounted for 61.1% of energy imports and coal for 29.8%. Heavy Fuel Oil (HFO) powers "
        "the CEB's base-load power plants, while kerosene drives gas turbines during peak periods. "
        "The sugar industry contributes bagasse for cogeneration, providing approximately 10% of "
        "electricity generation. Renewable energy (solar, wind, hydro) contributed 18.2% to the "
        "electricity mix in 2024."
    )

    doc.add_heading("11.2 Petroleum Pricing Mechanism (PPM)", level=2)
    doc.add_paragraph(
        "The PPM, administered by the State Trading Corporation and overseen by the Petroleum Pricing "
        "Committee, determines retail fuel prices for Mogas (RON 95) and Gas Oil (diesel). Prices are "
        "computed based on a 6-month average Platts price (3 months actual + 3 months future) plus "
        "a margin of up to 4%. The Price Stabilisation Account (PSA) buffers windfall gains and losses, "
        "preventing retail price volatility. As of April 2026, Mogas retails at Rs 64.25/L and Gas Oil "
        "at Rs 71.25/L."
    )

    doc.add_heading("11.3 Renewable Energy Transition", level=2)
    doc.add_paragraph(
        "Budget 2025-2026 committed MUR 30 billion to solar energy and biomass projects, targeting "
        "35% renewable energy in the electricity mix by 2028 and 60% by 2030. MARENA (Mauritius "
        "Renewable Energy Agency) coordinates deployment of 300,000 rooftop solar systems, 150 MW "
        "in agro-industrial solar, and 30 hybrid PV-BESS systems. The EV transition is underway with "
        "electric buses being added to the public transport fleet."
    )

    doc.add_heading("11.4 Carbon Reduction Commitments", level=2)
    doc.add_paragraph(
        "Under the Paris Agreement, Mauritius committed to reducing CO2 emissions by 4.3 million tons. "
        "The AstraFlow carbon footprint calculator aligns with this target by helping sector stakeholders "
        "quantify their fuel-related emissions and identify reduction opportunities through fuel switching, "
        "efficiency improvements, and renewable energy adoption."
    )

    doc.add_page_break()

    # ═══════════════════ 12. FUTURE ROADMAP ══════════════════════════════
    doc.add_heading("12. Future Roadmap", level=1)

    doc.add_heading("12.1 ML/AI Predictive Modelling for Mauritius", level=2)
    doc.add_paragraph(
        "The backend/ml/ directory and installed scikit-learn, pandas, and numpy dependencies position "
        "AstraFlow for future ML capabilities. Planned models include:"
    )
    for item in [
        "Time-series forecasting of STC fuel prices using historical PPM adjustment data and Brent crude correlations.",
        "Predictive modelling of CEB electricity demand based on sector-specific fuel consumption patterns.",
        "Bagasse yield prediction for sugar mills using climate and harvest data.",
        "Fleet fuel consumption optimisation using route and load pattern analysis.",
        "Automated impact assessment for PPM threshold changes on each Mauritius sector.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("12.2 Planned Enhancements", level=2)
    for item in [
        "Integration with STC's real-time fuel price API for live PPM data.",
        "CEB tariff API integration for combined fuel + electricity cost analysis.",
        "MARENA renewable energy subsidy calculator for solar/biomass project ROI.",
        "Push notifications for PPC price adjustment announcements.",
        "Secure mobile auth via biometrics for government-issued devices.",
        "Offline-first data persistence for field use in remote areas.",
        "Role-based access control for multi-tenant government deployment.",
        "Asymmetric JWT signing (RS256) for distributed CEB/Ministry architecture.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ── Footer ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Document —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("747683")
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AstraFlow — Empowering Mauritius' Energy Transition")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("747683")

    return doc


if __name__ == "__main__":
    doc = build_document()
    output = "AstraFlow_System_Design_Architecture_Mauritius.docx"
    doc.save(output)
    print(f"Document generated: {output}")
