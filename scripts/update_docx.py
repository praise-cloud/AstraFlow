"""Append the forecasting model implementation section to the .docx."""
from docx import Document
from docx.shared import Pt

doc = Document("AstraFlow_Issues_and_Resolutions.docx")

# ---- New section: Forecasting Model Implementation ----
doc.add_heading("Forecasting Model — Implementation, Execution & Continuous Training", level=1)

# -- Implementation Overview --
doc.add_heading("Implementation", level=2)

p = doc.add_paragraph()
p.add_run("The forecasting model follows a 5-stage pipeline implemented across three files:")

# Pipeline table
table = doc.add_table(rows=6, cols=3, style="Light Grid Accent 1")
headers = ["Stage", "File", "Function / Class"]
rows_data = [
    ["1. Load Historical Data", "backend/ml/data.py", "load_training_data_from_db() / generate_training_data()"],
    ["2. Clean Data", "backend/ml/data.py", "clean_data()"],
    ["3. Feature Extraction", "backend/ml/data.py", "prepare_features() / build_arima_series()"],
    ["4. Model Training", "backend/ml/models.py", "EnsembleForecaster.fit() - trains Linear + RF + ARIMA"],
    ["5. Forecast Generation", "backend/ml/forecast.py", "FuelForecaster.forecast()"],
]
for i, row_data in enumerate([headers] + rows_data):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.bold = True

doc.add_paragraph()

# Three model types
doc.add_heading("Three Model Types", level=3)

p = doc.add_paragraph()
r = p.add_run("1. Linear Regression")
r.bold = True
p.add_run(" via sklearn.linear_model.LinearRegression. Falls back to pure NumPy OLS (Ordinary Least Squares) when sklearn is unavailable. Uses the 6 engineered features (day_index, weekly_sin, weekly_cos, yearly_sin, momentum, weekly_diff).")

p = doc.add_paragraph()
r = p.add_run("2. Random Forest")
r.bold = True
p.add_run(" via sklearn.ensemble.RandomForestRegressor (100 trees, max_depth=6). Trained on the same feature matrix as Linear Regression. Provides feature importance rankings.")

p = doc.add_paragraph()
r = p.add_run("3. ARIMA Time Series")
r.bold = True
p.add_run(" via statsmodels.tsa.arima.model.ARIMA with order (2,1,2). Operates directly on the raw 1-D price series, not engineered features. Captures autocorrelation and trend patterns that regression models may miss.")

p = doc.add_paragraph()
r = p.add_run("Ensemble blending: ")
r.bold = True
p.add_run("The final forecast is a weighted average of the regression ensemble (Linear + RF) and the ARIMA forecast. The default weight is 70% regression / 30% ARIMA. If ARIMA is unavailable (statsmodels not installed), pure regression is used.")

# -- Stage details --
doc.add_heading("Stage Details", level=3)

p = doc.add_paragraph(style="List Bullet")
r = p.add_run("Stage 1 \u2014 Historical Data")
r.bold = True
doc.add_paragraph("Loads up to 365 days of fuel price records from the fuel_prices table. If the database is empty or contains fewer than 14 records, synthetic training data is generated programmatically with realistic seasonal patterns (yearly sin wave, weekly oscillation, Gaussian noise).")

p = doc.add_paragraph(style="List Bullet")
r = p.add_run("Stage 2 \u2014 Data Cleaning")
r.bold = True
doc.add_paragraph("clean_data() performs four operations in sequence: (a) sorts chronologically, (b) removes rows where the target fuel type is None, (c) removes statistical outliers via the IQR method (1.5x IQR beyond Q1/Q3), (d) forward-fills any missing dates within the date range using the last known price.")

p = doc.add_paragraph(style="List Bullet")
r = p.add_run("Stage 3 \u2014 Feature Extraction")
r.bold = True
doc.add_paragraph("prepare_features() builds a 6-column feature matrix: day_index (linear day counter), weekly_sin/cos (sin/cos of 2*pi*day/7 for weekly seasonality), yearly_sin (sin of 2*pi*day/365 for yearly seasonality), momentum (1-day lag diff), weekly_diff (7-day lag diff). For ARIMA, build_arima_series() returns the raw 1-D price array.")

p = doc.add_paragraph(style="List Bullet")
r = p.add_run("Stage 4 \u2014 Model Training")
r.bold = True
doc.add_paragraph("EnsembleForecaster.fit() trains all three models on the same historical window. Linear and RF use the feature matrix X. ARIMA fits directly on the raw price series y_raw. Each model availability flag is set individually.")

p = doc.add_paragraph(style="List Bullet")
r = p.add_run("Stage 5 \u2014 Forecast Generation")
r.bold = True
doc.add_paragraph("FuelForecaster.forecast() builds a future feature matrix with zero momentum/diff, generates blended predictions, computes 95% confidence intervals via prediction +/- 1.96*sigma of training residuals, classifies trend (up/down/stable based on 1% threshold), and produces a recommendation with urgency level.")

# -- Execution --
doc.add_heading("Execution", level=2)

p = doc.add_paragraph()
r = p.add_run("Server startup: ")
r.bold = True
p.add_run("When the FastAPI application starts, the lifespan handler in main.py calls init_db() and seed(). The first API request to any ML endpoint (dashboard, predict, forecast) triggers get_forecaster(), which creates the singleton FuelForecaster instance and calls train() - executing all 5 pipeline stages at once.")

p = doc.add_paragraph()
r = p.add_run("Per-request flow: ")
r.bold = True
p.add_run("The dashboard endpoint calls get_forecaster() which checks if retraining is needed (see below). The predict and forecast endpoints call FuelForecaster.forecast() directly. Each forecast call re-loads and re-cleans the latest data to ensure freshness, but model retraining only happens when new data appears.")

p = doc.add_paragraph()
r = p.add_run("Response shape: ")
r.bold = True
p.add_run("The forecast response includes: current_price, avg_forecast, min/max, trend direction and change percentage, 30-day point predictions with lower/upper bounds, confidence interval, model name, evaluation metrics (MAE, RMSE, R2, ARIMA AIC), feature importance rankings, and a recommendation with action/urgency.")

# -- Continuous training --
doc.add_heading("Continuous Training Strategy", level=2)

p = doc.add_paragraph()
r = p.add_run("The model does NOT retrain on every request. ")
r.bold = True
p.add_run("Three mechanisms keep the model current with minimal overhead:")

b1 = doc.add_paragraph(style="List Bullet")
r = b1.add_run("DB-change detection (lazy retrain): ")
r.bold = True
b1.add_run("Before each forecast, _retrain_if_needed() queries the database for the most recent price date. If the latest DB date is newer than the last training date, the model retrains automatically. This catches new scraped prices without manual triggers.")

b2 = doc.add_paragraph(style="List Bullet")
r = b2.add_run("Time-based refresh: ")
r.bold = True
b2.add_run("The singleton accessor get_forecaster() also forces a retrain every 6 hours (configured via _RETRAIN_INTERVAL) regardless of whether new data was detected. This ensures the model does not go stale during long-running server instances.")

b3 = doc.add_paragraph(style="List Bullet")
r = b3.add_run("Cold start: ")
r.bold = True
b3.add_run("On first call (server restart or first request), the forecaster is None, so get_forecaster() creates a new instance and calls train(). The synthetic data fallback ensures the model always has at least 365 training points to work with, even on a fresh database.")

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run("Future considerations: ")
r.bold = True
p.add_run("For production, consider persisting trained model coefficients to disk (joblib/pickle) so the model survives server restarts without retraining. The current implementation retrains from scratch on each cold start, which takes approximately 1-2 seconds for 365 data points.")

doc.save("AstraFlow_Issues_and_Resolutions.docx")
print("Done - document updated")
